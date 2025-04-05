#!/usr/bin/env python3
"""
Service implementation for Document Management operations.
"""
import os
import json
import logging
import tempfile
import io
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
import mimetypes

from app.tools.base.service import ToolServiceBase


class DocumentManagementService(ToolServiceBase):
    """Service to handle document management operations"""

    SERVICE_NAME = "document_management"
    REQUIRED_ENV_VARS = []
    OPTIONAL_ENV_VARS = {
        "DOCUMENT_CACHE_DIR": "/tmp/document_cache",
        "MAX_DOCUMENT_SIZE_MB": "50",
        "ALLOWED_DOCUMENT_TYPES": "pdf,docx,txt,md,html,csv,xlsx,pptx,json,xml",
        "RATE_LIMIT_DOCUMENTS": "100/hour"
    }

    def __init__(self):
        """Initialize the Document Management service"""
        super().__init__()
        self.cache_dir = self.get_env_var(
            "DOCUMENT_CACHE_DIR", default="/tmp/document_cache")
        self.max_size_mb = int(self.get_env_var(
            "MAX_DOCUMENT_SIZE_MB", default="50"))
        self.allowed_types = self.get_env_var("ALLOWED_DOCUMENT_TYPES",
                                              default="pdf,docx,txt,md,html,csv,xlsx,pptx,json,xml").split(",")

        # Initialize parsers and converters as empty dictionaries
        self.parsers = {}
        self.converters = {}

        # Set up rate limiting
        rate_limit_spec = self.get_env_var(
            "RATE_LIMIT_DOCUMENTS", default="100/hour")
        if "/" in rate_limit_spec:
            calls, period = rate_limit_spec.split("/")
            calls = int(calls)
            if period.lower() == "second":
                period_seconds = 1
            elif period.lower() == "minute":
                period_seconds = 60
            elif period.lower() == "hour":
                period_seconds = 3600
            elif period.lower() == "day":
                period_seconds = 86400
            else:
                period_seconds = 3600  # Default to hour

            self.create_rate_limiter(
                "document_operations", calls, period_seconds)

    def initialize(self) -> bool:
        """
        Initialize the Document Management service.

        Returns:
            bool: True if initialization was successful, False otherwise
        """
        try:
            # Ensure cache directory exists
            os.makedirs(self.cache_dir, exist_ok=True)

            # Initialize document parsers
            self._load_parsers()

            # Initialize document converters
            self._load_converters()

            self.initialized = True
            self.logger.info(
                "Document Management service initialized successfully")
            return True
        except Exception as e:
            self.logger.error(
                f"Failed to initialize Document Management service: {str(e)}")
            return False

    def _load_parsers(self):
        """Load available document parsers based on installed dependencies"""
        parsers = {}

        # Text-based document parsers
        try:
            parsers["txt"] = self._parse_text
            parsers["md"] = self._parse_text
            parsers["html"] = self._parse_text
            parsers["xml"] = self._parse_text
            parsers["json"] = self._parse_json
        except ImportError:
            self.logger.warning("Basic text parser registration failed")

        # DOCX parser
        try:
            import docx
            parsers["docx"] = self._parse_docx
        except ImportError:
            self.logger.warning(
                "DOCX parser not available - python-docx package missing")

        # PDF parser (using PyPDF2)
        try:
            import PyPDF2
            parsers["pdf"] = self._parse_pdf
        except ImportError:
            self.logger.warning(
                "PDF parser not available - PyPDF2 package missing")

        # Excel parser
        try:
            import pandas as pd
            parsers["xlsx"] = self._parse_excel
            parsers["csv"] = self._parse_csv
        except ImportError:
            self.logger.warning(
                "Excel/CSV parser not available - pandas package missing")

        # Update the parsers dictionary
        self.parsers = parsers
        self.logger.info(
            f"Loaded document parsers for: {', '.join(parsers.keys())}")

    def _load_converters(self):
        """Load available document converters based on installed dependencies"""
        converters = {}

        # Text format converters
        try:
            converters["txt_to_html"] = self._convert_text_to_html
            converters["html_to_txt"] = self._convert_html_to_text
            converters["md_to_html"] = self._convert_markdown_to_html
            converters["html_to_md"] = self._convert_html_to_markdown
        except ImportError:
            self.logger.warning("Basic text converters registration failed")

        # Try to load pandoc-based converters if available
        try:
            import pypandoc
            converters["docx_to_html"] = self._convert_docx_to_html
            converters["docx_to_md"] = self._convert_docx_to_markdown
            converters["html_to_docx"] = self._convert_html_to_docx
            converters["md_to_docx"] = self._convert_markdown_to_docx
        except ImportError:
            self.logger.warning(
                "Pandoc converters not available - pypandoc package missing")

        # Update the converters dictionary
        self.converters = converters
        self.logger.info(
            f"Loaded document converters for: {', '.join(converters.keys())}")

    def validate_document(self, document_path: str) -> Dict[str, Any]:
        """
        Validate document size and type.

        Args:
            document_path: Path to the document

        Returns:
            Dict with validation result
        """
        self._is_initialized()

        try:
            # Check if file exists
            if not os.path.exists(document_path):
                return {
                    "status": "error",
                    "message": f"Document not found: {document_path}"
                }

            # Check if it's a file
            if not os.path.isfile(document_path):
                return {
                    "status": "error",
                    "message": f"Not a file: {document_path}"
                }

            # Get file extension and check if it's allowed
            _, ext = os.path.splitext(document_path)
            ext = ext.lower().lstrip(".")

            if ext not in self.allowed_types:
                return {
                    "status": "error",
                    "message": f"Document type not allowed: {ext}. Allowed types: {', '.join(self.allowed_types)}"
                }

            # Check file size
            size_bytes = os.path.getsize(document_path)
            size_mb = size_bytes / (1024 * 1024)

            if size_mb > self.max_size_mb:
                return {
                    "status": "error",
                    "message": f"Document too large: {size_mb:.2f} MB. Maximum allowed size: {self.max_size_mb} MB"
                }

            # Determine MIME type
            mime_type, _ = mimetypes.guess_type(document_path)

            return {
                "status": "success",
                "info": {
                    "path": document_path,
                    "filename": os.path.basename(document_path),
                    "extension": ext,
                    "mime_type": mime_type,
                    "size_bytes": size_bytes,
                    "size_mb": size_mb
                }
            }
        except Exception as e:
            self.logger.error(f"Error validating document: {str(e)}")
            return {
                "status": "error",
                "message": f"Error validating document: {str(e)}"
            }

    def get_document_info(self, document_path: str) -> Dict[str, Any]:
        """
        Get metadata information about a document.

        Args:
            document_path: Path to the document

        Returns:
            Dict with document metadata
        """
        self._is_initialized()

        # First validate the document
        validation = self.validate_document(document_path)
        if validation["status"] == "error":
            return validation

        # Get basic info from validation
        info = validation["info"]
        ext = info["extension"]

        try:
            # Get additional metadata based on file type
            if ext == "pdf":
                # Get PDF metadata using PyPDF2
                try:
                    import PyPDF2
                    with open(document_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        info["pages"] = len(reader.pages)
                        info["metadata"] = {}
                        if reader.metadata:
                            for key in reader.metadata:
                                value = reader.metadata[key]
                                if isinstance(value, bytes):
                                    try:
                                        value = value.decode('utf-8')
                                    except UnicodeDecodeError:
                                        value = str(value)
                                info["metadata"][key] = value
                        info["is_encrypted"] = reader.is_encrypted
                except ImportError:
                    info["metadata_error"] = "PyPDF2 not available for PDF analysis"

            elif ext == "docx":
                # Get DOCX metadata using python-docx
                try:
                    import docx
                    doc = docx.Document(document_path)
                    info["metadata"] = {}
                    for prop in doc.core_properties.__dict__:
                        if not prop.startswith('_') and hasattr(doc.core_properties, prop):
                            value = getattr(doc.core_properties, prop)
                            if value is not None:
                                info["metadata"][prop] = str(value)

                    # Count paragraphs, tables, etc.
                    info["paragraphs"] = len(doc.paragraphs)
                    info["tables"] = len(doc.tables)
                    info["sections"] = len(doc.sections)
                except ImportError:
                    info["metadata_error"] = "python-docx not available for DOCX analysis"

            elif ext in ["xlsx", "csv"]:
                # Get spreadsheet metadata using pandas
                try:
                    import pandas as pd
                    if ext == "xlsx":
                        # Get Excel workbook sheet names
                        info["sheet_names"] = pd.ExcelFile(
                            document_path).sheet_names
                        # Get row count for first sheet
                        first_sheet = info["sheet_names"][0]
                        df = pd.read_excel(
                            document_path, sheet_name=first_sheet)
                        info["rows"] = len(df)
                        info["columns"] = len(df.columns)
                    else:  # CSV
                        df = pd.read_csv(document_path)
                        info["rows"] = len(df)
                        info["columns"] = len(df.columns)
                except ImportError:
                    info["metadata_error"] = "pandas not available for spreadsheet analysis"

            # Basic text file info
            elif ext in ["txt", "md", "html", "xml", "json"]:
                with open(document_path, 'r', encoding='utf-8', errors='replace') as file:
                    content = file.read()
                    info["lines"] = len(content.splitlines())
                    info["characters"] = len(content)

                    # For JSON, validate and add structure info
                    if ext == "json":
                        try:
                            json_data = json.loads(content)
                            if isinstance(json_data, dict):
                                info["json_keys"] = list(json_data.keys())
                            elif isinstance(json_data, list):
                                info["json_items"] = len(json_data)
                        except json.JSONDecodeError:
                            info["json_error"] = "Invalid JSON format"

            return {
                "status": "success",
                "info": info
            }
        except Exception as e:
            self.logger.error(f"Error getting document info: {str(e)}")
            return {
                "status": "error",
                "message": f"Error getting document info: {str(e)}"
            }

    def _parse_text(self, document_path: str) -> Dict[str, Any]:
        """Parse text-based document"""
        with open(document_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            lines = content.splitlines()

            return {
                "text": content,
                "lines": lines,
                "line_count": len(lines),
                "character_count": len(content)
            }

    def _parse_json(self, document_path: str) -> Dict[str, Any]:
        """Parse JSON document"""
        with open(document_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
            data = json.loads(content)

            return {
                "data": data,
                "structure": self._analyze_json_structure(data)
            }

    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "count": len(data)
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "count": len(data),
                "sample": data[:5] if len(data) > 0 else []
            }
        else:
            return {
                "type": type(data).__name__
            }

    def _parse_docx(self, document_path: str) -> Dict[str, Any]:
        """Parse DOCX document"""
        import docx
        doc = docx.Document(document_path)

        # Extract paragraphs
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append({
                    "text": p.text,
                    "style": p.style.name
                })

        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                "id": i,
                "rows": len(table.rows),
                "columns": len(table.rows[0].cells) if table.rows else 0,
                "data": table_data
            })

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables)
        }

    def _parse_pdf(self, document_path: str) -> Dict[str, Any]:
        """Parse PDF document"""
        import PyPDF2
        with open(document_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)

            # Extract text from each page
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                pages.append({
                    "page_number": i + 1,
                    "text": text
                })

            return {
                "pages": pages,
                "page_count": len(pages)
            }

    def _parse_excel(self, document_path: str) -> Dict[str, Any]:
        """Parse Excel document"""
        import pandas as pd

        # Get all sheet names
        xlsx = pd.ExcelFile(document_path)
        sheets = xlsx.sheet_names

        result = {
            "sheets": {},
            "sheet_names": sheets
        }

        # Read each sheet (limit to first 10 rows for preview)
        for sheet in sheets:
            df = pd.read_excel(document_path, sheet_name=sheet, nrows=10)
            result["sheets"][sheet] = {
                "columns": list(df.columns),
                "preview": df.to_dict(orient="records"),
                "total_rows": len(pd.read_excel(document_path, sheet_name=sheet))
            }

        return result

    def _parse_csv(self, document_path: str) -> Dict[str, Any]:
        """Parse CSV document"""
        import pandas as pd

        # Read CSV file
        df = pd.read_csv(document_path)

        return {
            "columns": list(df.columns),
            "preview": df.head(10).to_dict(orient="records"),
            "total_rows": len(df)
        }

    def parse_document(self, document_path: str) -> Dict[str, Any]:
        """
        Parse a document into structured content.

        Args:
            document_path: Path to the document

        Returns:
            Dict with parsed document content
        """
        self._is_initialized()

        # First validate the document
        validation = self.validate_document(document_path)
        if validation["status"] == "error":
            return validation

        # Get basic info from validation
        info = validation["info"]
        ext = info["extension"]

        # Apply rate limiting
        self.apply_rate_limit("document_operations", 100, 3600, wait=True)

        try:
            # Check if we have a parser for this extension
            if ext not in self.parsers:
                return {
                    "status": "error",
                    "message": f"No parser available for document type: {ext}"
                }

            # Call the appropriate parser function
            parser_func = self.parsers[ext]
            result = parser_func(document_path)

            return {
                "status": "success",
                "info": info,
                "content": result
            }
        except Exception as e:
            self.logger.error(f"Error parsing document: {str(e)}")
            return {
                "status": "error",
                "message": f"Error parsing document: {str(e)}"
            }

    def convert_document(self, input_path: str, output_path: str, output_format: str) -> Dict[str, Any]:
        """
        Convert a document from one format to another.

        Args:
            input_path: Path to the input document
            output_path: Path for the output document
            output_format: Target format

        Returns:
            Dict with conversion result
        """
        self._is_initialized()

        # First validate the input document
        validation = self.validate_document(input_path)
        if validation["status"] == "error":
            return validation

        # Get input format from validation
        input_info = validation["info"]
        input_format = input_info["extension"]

        # Check if output format is valid
        if output_format not in self.allowed_types:
            return {
                "status": "error",
                "message": f"Output format not allowed: {output_format}. Allowed formats: {', '.join(self.allowed_types)}"
            }

        # Determine converter key
        converter_key = f"{input_format}_to_{output_format}"

        # Apply rate limiting
        self.apply_rate_limit("document_operations", 100, 3600, wait=True)

        try:
            # Check if we have a converter for this combination
            if converter_key not in self.converters:
                return {
                    "status": "error",
                    "message": f"No converter available for {input_format} to {output_format}"
                }

            # Create output directory if needed
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)

            # Call the appropriate converter function
            converter_func = self.converters[converter_key]
            result = converter_func(input_path, output_path)

            return {
                "status": "success",
                "input": input_info,
                "output": {
                    "path": output_path,
                    "format": output_format
                },
                "result": result
            }
        except Exception as e:
            self.logger.error(f"Error converting document: {str(e)}")
            return {
                "status": "error",
                "message": f"Error converting document: {str(e)}"
            }

    def _convert_text_to_html(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert text file to HTML"""
        with open(input_path, 'r', encoding='utf-8', errors='replace') as in_file:
            text = in_file.read()

            # Simple text to HTML conversion
            html = f"<!DOCTYPE html>\n<html>\n<head>\n<title>{os.path.basename(input_path)}</title>\n</head>\n<body>\n"

            # Convert line breaks to <p> tags
            for line in text.split('\n'):
                if line.strip():
                    html += f"<p>{line}</p>\n"
                else:
                    html += "<br>\n"

            html += "</body>\n</html>"

            # Write output
            with open(output_path, 'w', encoding='utf-8') as out_file:
                out_file.write(html)

            return {
                "converted": True
            }

    def _convert_html_to_text(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert HTML to plain text"""
        try:
            from bs4 import BeautifulSoup

            with open(input_path, 'r', encoding='utf-8', errors='replace') as in_file:
                html = in_file.read()

                # Parse HTML
                soup = BeautifulSoup(html, 'html.parser')

                # Extract text
                text = soup.get_text(separator='\n')

                # Write output
                with open(output_path, 'w', encoding='utf-8') as out_file:
                    out_file.write(text)

                return {
                    "converted": True
                }
        except ImportError:
            return {
                "converted": False,
                "error": "BeautifulSoup module not available"
            }

    def _convert_markdown_to_html(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert Markdown to HTML"""
        try:
            import markdown

            with open(input_path, 'r', encoding='utf-8', errors='replace') as in_file:
                md_text = in_file.read()

                # Convert Markdown to HTML
                html = markdown.markdown(md_text)

                # Add HTML document wrapper
                full_html = f"<!DOCTYPE html>\n<html>\n<head>\n<title>{os.path.basename(input_path)}</title>\n</head>\n<body>\n{html}\n</body>\n</html>"

                # Write output
                with open(output_path, 'w', encoding='utf-8') as out_file:
                    out_file.write(full_html)

                return {
                    "converted": True
                }
        except ImportError:
            return {
                "converted": False,
                "error": "markdown module not available"
            }

    def _convert_html_to_markdown(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert HTML to Markdown"""
        try:
            import html2text

            with open(input_path, 'r', encoding='utf-8', errors='replace') as in_file:
                html = in_file.read()

                # Convert HTML to Markdown
                h = html2text.HTML2Text()
                h.ignore_links = False
                md_text = h.handle(html)

                # Write output
                with open(output_path, 'w', encoding='utf-8') as out_file:
                    out_file.write(md_text)

                return {
                    "converted": True
                }
        except ImportError:
            return {
                "converted": False,
                "error": "html2text module not available"
            }

    def _convert_docx_to_html(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """Convert DOCX to HTML using pandoc"""
        try:
            import pypandoc

            # Convert using pandoc
            pypandoc.convert_file(input_path, 'html', outputfile=output_path)

            return {
                "converted": True
            }
        except ImportError:
            return {
                "converted": False,
                "error": "pypandoc module not available"
            }

    def extract_text(self, document_path: str) -> Dict[str, Any]:
        """
        Extract text content from a document.

        Args:
            document_path: Path to the document

        Returns:
            Dict with extracted text
        """
        self._is_initialized()

        # First validate the document
        validation = self.validate_document(document_path)
        if validation["status"] == "error":
            return validation

        # Get input format from validation
        info = validation["info"]
        format = info["extension"]

        # Apply rate limiting
        self.apply_rate_limit("document_operations", 100, 3600, wait=True)

        try:
            # Handle different document types
            if format == "pdf":
                # Extract text from PDF
                import PyPDF2
                with open(document_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n\n"

                return {
                    "status": "success",
                    "text": text,
                    "pages": len(reader.pages)
                }

            elif format == "docx":
                # Extract text from DOCX
                import docx
                doc = docx.Document(document_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"

                return {
                    "status": "success",
                    "text": text,
                    "paragraphs": len(doc.paragraphs)
                }

            elif format in ["txt", "md", "html", "xml", "json"]:
                # Extract text from text files
                with open(document_path, 'r', encoding='utf-8', errors='replace') as file:
                    text = file.read()

                # For HTML, try to extract clean text
                if format == "html":
                    try:
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(text, 'html.parser')
                        text = soup.get_text(separator='\n')
                    except ImportError:
                        pass

                return {
                    "status": "success",
                    "text": text,
                    "characters": len(text),
                    "lines": len(text.splitlines())
                }

            elif format in ["xlsx", "csv"]:
                # Extract text from spreadsheet
                import pandas as pd
                if format == "xlsx":
                    # Read all sheets
                    xlsx = pd.ExcelFile(document_path)
                    sheets = {}
                    for sheet_name in xlsx.sheet_names:
                        df = pd.read_excel(
                            document_path, sheet_name=sheet_name)
                        sheets[sheet_name] = df.to_csv(index=False)

                    return {
                        "status": "success",
                        "sheets": sheets,
                        "sheet_count": len(sheets)
                    }
                else:  # CSV
                    df = pd.read_csv(document_path)
                    return {
                        "status": "success",
                        "text": df.to_csv(index=False),
                        "rows": len(df),
                        "columns": len(df.columns)
                    }

            else:
                return {
                    "status": "error",
                    "message": f"Text extraction not supported for format: {format}"
                }

        except Exception as e:
            self.logger.error(f"Error extracting text: {str(e)}")
            return {
                "status": "error",
                "message": f"Error extracting text: {str(e)}"
            }

    def search_document(self, document_path: str, query: str, case_sensitive: bool = False) -> Dict[str, Any]:
        """
        Search for text within a document.

        Args:
            document_path: Path to the document
            query: Text to search for
            case_sensitive: Whether the search is case-sensitive

        Returns:
            Dict with search results
        """
        self._is_initialized()

        # Extract text from the document
        text_result = self.extract_text(document_path)

        if text_result["status"] == "error":
            return text_result

        # Apply rate limiting
        self.apply_rate_limit("document_operations", 100, 3600, wait=True)

        try:
            import re

            results = []

            # Handle different document types
            if "text" in text_result:
                # Simple text search for document with single text field
                text = text_result["text"]

                # Prepare search pattern
                pattern = re.escape(query)
                flags = 0 if case_sensitive else re.IGNORECASE

                # Find all matches
                matches = list(re.finditer(pattern, text, flags))

                # Create context for each match
                for match in matches:
                    start, end = match.span()

                    # Get context (up to 50 chars before and after)
                    context_start = max(0, start - 50)
                    context_end = min(len(text), end + 50)

                    # Get line number
                    line_number = text[:start].count('\n') + 1

                    results.append({
                        "match": match.group(),
                        "position": start,
                        "line": line_number,
                        "context": text[context_start:context_end]
                    })

            elif "sheets" in text_result:
                # Search across multiple sheets
                for sheet_name, sheet_text in text_result["sheets"].items():
                    # Prepare search pattern
                    pattern = re.escape(query)
                    flags = 0 if case_sensitive else re.IGNORECASE

                    # Find all matches
                    matches = list(re.finditer(pattern, sheet_text, flags))

                    # Create context for each match
                    for match in matches:
                        start, end = match.span()

                        # Get context (up to 50 chars before and after)
                        context_start = max(0, start - 50)
                        context_end = min(len(sheet_text), end + 50)

                        # Get line number (approximate row in the sheet)
                        line_number = sheet_text[:start].count('\n') + 1

                        results.append({
                            "match": match.group(),
                            "position": start,
                            "sheet": sheet_name,
                            "row": line_number,
                            "context": sheet_text[context_start:context_end]
                        })

            return {
                "status": "success",
                "query": query,
                "case_sensitive": case_sensitive,
                "matches": len(results),
                "results": results
            }

        except Exception as e:
            self.logger.error(f"Error searching document: {str(e)}")
            return {
                "status": "error",
                "message": f"Error searching document: {str(e)}"
            }


# Singleton instance
_service_instance = None


def get_service() -> DocumentManagementService:
    """
    Get or initialize the Document Management service singleton.

    Returns:
        DocumentManagementService instance
    """
    global _service_instance
    if _service_instance is None:
        _service_instance = DocumentManagementService()
        _service_instance.initialize()
    return _service_instance
